"""
Script pour générer un CV au format DOCX basé sur le design HTML/CSS
Auteur: Manuel Petis MFOU'OU
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_background_color(cell, color):
    """Ajouter une couleur de fond à une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_cv():
    # Créer le document
    doc = Document()
    
    # Définir les marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # COULEURS du design (en hexadécimal converti en RGB)
    primary_color = RGBColor(29, 50, 81)      # #1D3251
    secondary_color = RGBColor(205, 237, 218) # #CDEDDA
    text_color = RGBColor(38, 38, 38)         # #262626
    
    # ====== EN-TÊTE ======
    # Créer une table pour l'en-tête (3 colonnes)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Configurer les largeurs des colonnes
    header_table.columns[0].width = Cm(4)
    header_table.columns[1].width = Cm(10)
    header_table.columns[2].width = Cm(4)
    
    # Appliquer la couleur de fond primary_color
    for cell in header_table.rows[0].cells:
        add_background_color(cell, '1D3251')
    
    # Colonne 1: Sigle
    sigle_cell = header_table.rows[0].cells[0]
    sigle_para = sigle_cell.paragraphs[0]
    sigle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sigle_para.add_run('m/mp')
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Colonne 2: Nom et titre
    name_cell = header_table.rows[0].cells[1]
    name_cell.paragraphs[0].clear()
    
    # Nom de famille
    lastname_para = name_cell.add_paragraph()
    lastname_run = lastname_para.add_run("MFOU'OU")
    lastname_run.font.size = Pt(28)
    lastname_run.font.bold = True
    lastname_run.font.color.rgb = secondary_color
    
    # Prénom
    firstname_para = name_cell.add_paragraph()
    firstname_run = firstname_para.add_run('Manuel Petis')
    firstname_run.font.size = Pt(28)
    firstname_run.font.color.rgb = secondary_color
    
    # Titre
    title_para = name_cell.add_paragraph()
    title_run = title_para.add_run('Ingénieur Génie Civil')
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Colonne 3: Photo (placeholder)
    photo_cell = header_table.rows[0].cells[2]
    photo_para = photo_cell.paragraphs[0]
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_run = photo_para.add_run('[Photo]')
    photo_run.font.size = Pt(12)
    photo_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Barre secondaire
    doc.add_paragraph()
    bar_table = doc.add_table(rows=1, cols=1)
    bar_cell = bar_table.rows[0].cells[0]
    add_background_color(bar_cell, 'CDEDDA')
    bar_cell.paragraphs[0].add_run(' ')
    
    # ====== CORPS DU CV (2 colonnes) ======
    doc.add_paragraph()
    
    # Table principale à 2 colonnes
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.columns[0].width = Cm(7)
    main_table.columns[1].width = Cm(11)
    
    left_cell = main_table.rows[0].cells[0]
    right_cell = main_table.rows[0].cells[1]
    
    # ====== COLONNE GAUCHE ======
    left_cell.paragraphs[0].clear()
    
    # CONTACT
    contact_heading = left_cell.add_paragraph()
    contact_heading_run = contact_heading.add_run('CONTACT')
    contact_heading_run.font.size = Pt(14)
    contact_heading_run.font.bold = True
    contact_heading_run.font.color.rgb = primary_color
    
    contact_items = [
        ('📞', '+237 697115067'),
        ('📧', 'manoumanuel19gmail.com'),
        ('📍', 'Yaoundé, Cameroun')
    ]
    
    for icon, text in contact_items:
        p = left_cell.add_paragraph()
        run = p.add_run(f'{icon} {text}')
        run.font.size = Pt(10)
        run.font.color.rgb = text_color
    
    left_cell.add_paragraph()
    
    # LANGUES
    lang_heading = left_cell.add_paragraph()
    lang_heading_run = lang_heading.add_run('LANGUES')
    lang_heading_run.font.size = Pt(14)
    lang_heading_run.font.bold = True
    lang_heading_run.font.color.rgb = primary_color
    
    languages = [
        ('🇫🇷', 'Français', 'Courant'),
        ('🇬🇧', 'Anglais', 'Intermédiaire')
    ]
    
    for flag, name, level in languages:
        p = left_cell.add_paragraph()
        run = p.add_run(f'{flag} {name}: {level}')
        run.font.size = Pt(10)
        run.font.color.rgb = text_color
    
    left_cell.add_paragraph()
    
    # COMPÉTENCES
    skills_heading = left_cell.add_paragraph()
    skills_heading_run = skills_heading.add_run('SKILLS')
    skills_heading_run.font.size = Pt(14)
    skills_heading_run.font.bold = True
    skills_heading_run.font.color.rgb = primary_color
    
    skill_categories = {
        'Logiciels de CAO/DAO': ['Revit', 'ArchiCAD', 'Robot Structural Analysis', 'Cype', 'Lumion'],
        'Programmation et Calculs': ['MATLAB', 'Python'],
        'Normes et Réglementations': ['Eurocode 2', 'BAEL 99'],
        'Modélisation de données': ['Power BI'],
        'Pack Office': ['Microsoft Word', 'Excel']
    }
    
    for category, skills in skill_categories.items():
        cat_para = left_cell.add_paragraph()
        cat_run = cat_para.add_run(category)
        cat_run.font.size = Pt(10)
        cat_run.font.bold = True
        cat_run.font.color.rgb = text_color
        
        for skill in skills:
            skill_para = left_cell.add_paragraph(style='List Bullet')
            skill_run = skill_para.add_run(f'{skill} (70%)')
            skill_run.font.size = Pt(9)
            skill_run.font.color.rgb = text_color
    
    left_cell.add_paragraph()
    
    # ÉDUCATION
    edu_heading = left_cell.add_paragraph()
    edu_heading_run = edu_heading.add_run('EDUCATION')
    edu_heading_run.font.size = Pt(14)
    edu_heading_run.font.bold = True
    edu_heading_run.font.color.rgb = primary_color
    
    education_items = [
        ('Cycle Ingénieur, Génie Civil', 'Ecole Nationale Supérieure Polytechnique de Yaoundé 1', '2021-2025'),
        ('Licence en Mathématiques', 'Université de Yaoundé 1', '2021'),
        ('Baccalauréat scientifique, Série C', 'Lycée Classique et Moderne de Sangmélima', '2018'),
        ('Probatoire scientifique, Série C', 'Lycée Classique et Moderne de Sangmélima', '2017'),
        ('BEPC', 'Lycée Classique et Moderne de Sangmélima', '2015')
    ]
    
    for diploma, school, year in education_items:
        diploma_para = left_cell.add_paragraph()
        diploma_run = diploma_para.add_run(diploma)
        diploma_run.font.size = Pt(10)
        diploma_run.font.bold = True
        diploma_run.font.color.rgb = text_color
        
        school_para = left_cell.add_paragraph()
        school_run = school_para.add_run(school)
        school_run.font.size = Pt(9)
        school_run.font.italic = True
        school_run.font.color.rgb = text_color
        
        year_para = left_cell.add_paragraph()
        year_run = year_para.add_run(year)
        year_run.font.size = Pt(9)
        year_run.font.color.rgb = text_color
        
        left_cell.add_paragraph()
    
    # ====== COLONNE DROITE ======
    right_cell.paragraphs[0].clear()
    
    # ABOUT ME
    about_heading = right_cell.add_paragraph()
    about_heading_run = about_heading.add_run('ABOUT ME')
    about_heading_run.font.size = Pt(14)
    about_heading_run.font.bold = True
    about_heading_run.font.color.rgb = primary_color
    
    about_para = right_cell.add_paragraph()
    about_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    about_run = about_para.add_run(
        "Diplômé de l'Ecole Nationale Supérieure Polytechnique de Yaoundé 1 en filière Génie Civil, "
        "passionné de conception et dimensionnement d'ouvrages."
    )
    about_run.font.size = Pt(11)
    about_run.font.color.rgb = text_color
    
    right_cell.add_paragraph()
    right_cell.add_paragraph()
    
    # EXPÉRIENCES
    exp_heading = right_cell.add_paragraph()
    exp_heading_run = exp_heading.add_run('EXPERIENCES')
    exp_heading_run.font.size = Pt(14)
    exp_heading_run.font.bold = True
    exp_heading_run.font.color.rgb = primary_color
    
    # Expérience 1
    exp1_year = right_cell.add_paragraph()
    exp1_year_run = exp1_year.add_run('Novembre 2025-Décembre 2025')
    exp1_year_run.font.size = Pt(10)
    exp1_year_run.font.bold = True
    exp1_year_run.font.color.rgb = text_color
    
    exp1_title = right_cell.add_paragraph()
    exp1_title_run = exp1_title.add_run('Stage Professionnel')
    exp1_title_run.font.size = Pt(11)
    exp1_title_run.font.bold = True
    exp1_title_run.font.color.rgb = text_color
    
    exp1_company = right_cell.add_paragraph()
    exp1_company_run = exp1_company.add_run('Le Cieu SARL - Yaoundé')
    exp1_company_run.font.size = Pt(10)
    exp1_company_run.font.italic = True
    exp1_company_run.font.bold = True
    exp1_company_run.font.color.rgb = text_color
    
    exp1_task = right_cell.add_paragraph(style='List Bullet')
    exp1_task_run = exp1_task.add_run("Suivi des travaux de finition du projet de construction de l'hôtel LA CONCORDE")
    exp1_task_run.font.size = Pt(10)
    exp1_task_run.font.color.rgb = text_color
    
    right_cell.add_paragraph()
    
    # Expérience 2
    exp2_year = right_cell.add_paragraph()
    exp2_year_run = exp2_year.add_run('Février 2025-Septembre 2025')
    exp2_year_run.font.size = Pt(10)
    exp2_year_run.font.bold = True
    exp2_year_run.font.color.rgb = text_color
    
    exp2_title = right_cell.add_paragraph()
    exp2_title_run = exp2_title.add_run('Stage ingénieur')
    exp2_title_run.font.size = Pt(11)
    exp2_title_run.font.bold = True
    exp2_title_run.font.color.rgb = text_color
    
    exp2_company = right_cell.add_paragraph()
    exp2_company_run = exp2_company.add_run('ERA-CAMEROUN - Yaoundé')
    exp2_company_run.font.size = Pt(10)
    exp2_company_run.font.italic = True
    exp2_company_run.font.bold = True
    exp2_company_run.font.color.rgb = text_color
    
    exp2_task = right_cell.add_paragraph(style='List Bullet')
    exp2_task_run = exp2_task.add_run('Suivi des prestations de gestion des déchets dans la ville de Yaoundé')
    exp2_task_run.font.size = Pt(10)
    exp2_task_run.font.color.rgb = text_color
    
    right_cell.add_paragraph()
    
    # Expérience 3
    exp3_year = right_cell.add_paragraph()
    exp3_year_run = exp3_year.add_run('Mai 2024-Septembre 2024')
    exp3_year_run.font.size = Pt(10)
    exp3_year_run.font.bold = True
    exp3_year_run.font.color.rgb = text_color
    
    exp3_title = right_cell.add_paragraph()
    exp3_title_run = exp3_title.add_run('Stage Pré-ingénieur')
    exp3_title_run.font.size = Pt(11)
    exp3_title_run.font.bold = True
    exp3_title_run.font.color.rgb = text_color
    
    exp3_company = right_cell.add_paragraph()
    exp3_company_run = exp3_company.add_run('HEAVY DUTY SARL Engineering - Yaoundé')
    exp3_company_run.font.size = Pt(10)
    exp3_company_run.font.italic = True
    exp3_company_run.font.bold = True
    exp3_company_run.font.color.rgb = text_color
    
    tasks_exp3 = [
        'Conception et modélisation 3D sous Revit et ArchiCAD',
        'Dimensionnement des élément structuraux sous Robot Structural Analysis',
        "Suivi des travaux de réalisation des éléments structuraux (gros œuvres) pour la construction de l'Hôtel CRYTAL PALACE"
    ]
    
    for task in tasks_exp3:
        task_para = right_cell.add_paragraph(style='List Bullet')
        task_run = task_para.add_run(task)
        task_run.font.size = Pt(10)
        task_run.font.color.rgb = text_color
    
    # ====== FOOTER ======
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© Cyrano Conseil 2025 - Tous droits réservés. Manuel Petis MFOU\'OU\'s Profile')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = primary_color
    
    # Sauvegarder le document
    output_path = 'public/doc/cv.docx'
    doc.save(output_path)
    print(f'✓ CV généré avec succès: {output_path}')

if __name__ == '__main__':
    create_cv()
